import os
import re
import csv
import json
import hmac
import hashlib
from datetime import datetime

import fitz  # PyMuPDF
import streamlit as st
import pandas as pd
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ------------------------------------------------------------------
# Config
# ------------------------------------------------------------------
BOT_NAME = "GurucoolBOT"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, "documents")
CACHE_FILE = os.path.join(BASE_DIR, "processed_data.json")
QA_LOG_FILE = os.path.join(BASE_DIR, "qa_log.csv")            # every question+answer gets appended here
SESSION_LOG_FILE = os.path.join(BASE_DIR, "session_log.csv")  # one row per login (now tracks real usernames)
USERS_FILE = os.path.join(BASE_DIR, "users.json")             # hashed+salted credentials

MIN_PARA_LEN = 40          # ignore tiny fragments (page numbers, headers)
TOP_CHUNKS = 8              # how many chunks to consider before sentence-ranking
TOP_MATCHES = 3             # how many distinct matches to show the user
MAX_ANSWER_SENTENCES = 2    # keep each match's snippet short
DOCX_SECTION_CHUNK_CHARS = 500  # roughly how much text goes in one docx chunk
CACHE_VERSION = 4           # bump whenever the chunk/metadata schema changes

SUPPORTED_EXTENSIONS = (".pdf", ".docx")


# ------------------------------------------------------------------
# 1. Document discovery + fingerprinting (so we know when to re-process)
# ------------------------------------------------------------------
def list_documents(folder):
    if not os.path.isdir(folder):
        return []
    return sorted(f for f in os.listdir(folder) if f.lower().endswith(SUPPORTED_EXTENSIONS))


def folder_fingerprint(folder):
    """A cheap hash of filenames + sizes + mtimes, so we can detect changes
    without re-reading every file's contents."""
    items = []
    for f in list_documents(folder):
        path = os.path.join(folder, f)
        stat = os.stat(path)
        items.append(f"{f}:{stat.st_size}:{int(stat.st_mtime)}")
    raw = "|".join(items)
    return hashlib.md5(raw.encode("utf-8")).hexdigest()


# ------------------------------------------------------------------
# 2a. PDF extraction — paragraphs tagged with "Page N"
# ------------------------------------------------------------------
AUTHOR_PATTERNS = [
    re.compile(r"\bby\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){0,3})"),
    re.compile(r"\bauthor[:\s]+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){0,3})", re.IGNORECASE),
    re.compile(r"\bwritten\s+by\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){0,3})", re.IGNORECASE),
]


def guess_author_from_text(pages_text):
    for text in pages_text:
        for pattern in AUTHOR_PATTERNS:
            match = pattern.search(text)
            if match:
                return match.group(1).strip()
    return None


def _rects_overlap(a, b):
    """a, b are (x0, y0, x1, y1) tuples."""
    return not (a[2] <= b[0] or b[2] <= a[0] or a[3] <= b[1] or b[3] <= a[1])


def extract_tables(page):
    """Detect tables on a page and return them as structured data:
    [{"bbox": (x0,y0,x1,y1), "header": [...], "rows": [[...], ...]}]
    Falls back to no tables found if the installed PyMuPDF is too old
    to support find_tables() (added in PyMuPDF 1.23)."""
    detected = []
    try:
        finder = page.find_tables()
    except AttributeError:
        return detected  # older PyMuPDF without table support — degrade gracefully
    for tbl in finder.tables:
        try:
            data = tbl.extract()
        except Exception:
            continue
        if not data or len(data) < 2:
            continue
        header = [(c or "").strip() for c in data[0]]
        rows = data[1:]
        detected.append({"bbox": tuple(tbl.bbox), "header": header, "rows": rows})
    return detected


def read_pdf(path, filename):
    """Returns (chunks, page_texts). Tables are detected and turned into one
    clean chunk per row (e.g. 'GL Code: 410020, Debit ($): 45,000.00, ...')
    instead of being flattened into a single unreadable text blob — that's
    what previously made a question like 'what is 410020' return the whole
    table jammed together instead of just that row."""
    chunks = []
    page_texts = {}
    doc = fitz.open(path)
    for page_no, page in enumerate(doc, start=1):
        location = f"Page {page_no}"
        page_texts[location] = page.get_text().strip()

        tables = extract_tables(page)
        table_bboxes = [t["bbox"] for t in tables]

        # 1. One clean chunk per table row, e.g.:
        #    "Region: USA | Invoice Nature Classification: Software
        #    Subscription & License | GL Code: 410020 | Debit ($): 45,000.00"
        for t_idx, table in enumerate(tables, start=1):
            header = table["header"]
            for r_idx, row in enumerate(table["rows"], start=1):
                cells = [(c or "").strip() for c in row]
                if not any(cells):
                    continue
                pairs = [
                    f"{(col or 'Value')}: {val}"
                    for col, val in zip(header, cells) if val
                ]
                if not pairs:
                    continue
                row_text = " | ".join(pairs)
                row_key = cells[0] if cells[0] else f"row {r_idx}"
                row_location = f"{location} — Table {t_idx}, Row {r_idx} ({row_key})"
                chunks.append({"source": filename, "location": row_location, "text": row_text})
                page_texts[row_location] = row_text

        # 2. Normal paragraph chunks for everything OUTSIDE detected tables,
        # so table text doesn't also get duplicated as a garbled blob.
        for block in page.get_text("blocks"):
            x0, y0, x1, y1, text = block[0], block[1], block[2], block[3], block[4]
            if any(_rects_overlap((x0, y0, x1, y1), tb) for tb in table_bboxes):
                continue
            para = re.sub(r"\s+", " ", text).strip()
            if len(para) > MIN_PARA_LEN:
                chunks.append({"source": filename, "location": location, "text": para})
    doc.close()
    return chunks, page_texts


def process_pdf(path, filename):
    doc = fitz.open(path)
    page_count = doc.page_count
    first_page_text = doc[0].get_text() if page_count else ""

    title = filename
    for line in first_page_text.split("\n"):
        line = line.strip()
        if len(line) > 4:
            title = line
            break

    meta_author = (doc.metadata or {}).get("author", "").strip()
    if meta_author:
        author = meta_author
    else:
        sample_pages = [doc[i].get_text() for i in range(min(2, page_count))]
        author = guess_author_from_text(sample_pages)
    doc.close()

    chunks, page_texts = read_pdf(path, filename)
    file_meta = {
        "filename": filename,
        "file_type": "pdf",
        "guessed_title": title,
        "author": author,
        "extent_label": f"{page_count} pages"
    }
    return chunks, page_texts, file_meta


# ------------------------------------------------------------------
# 2b. Word (.docx) extraction — paragraphs tagged with "Section: <heading>"
# ------------------------------------------------------------------
def read_docx(path, filename):
    """Returns (chunks, page_texts). Word has no fixed page numbers, so we
    cite by section (nearest preceding heading) instead — it's a stable
    reference the document itself defines."""
    document = Document(path)
    chunks = []
    sections = {}  # section title -> list of paragraph strings (for full-section reading)
    current_section = "Introduction"
    sections[current_section] = []
    buffer = []

    def flush_buffer():
        if buffer:
            text = " ".join(buffer).strip()
            if len(text) > MIN_PARA_LEN:
                chunks.append({
                    "source": filename,
                    "location": f"Section: {current_section}",
                    "text": text
                })
            buffer.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or style_name == "Title":
            flush_buffer()
            current_section = text
            sections.setdefault(current_section, [])
            continue
        sections[current_section].append(text)
        buffer.append(text)
        if sum(len(b) for b in buffer) > DOCX_SECTION_CHUNK_CHARS:
            flush_buffer()
    flush_buffer()

    page_texts = {
        f"Section: {title}": "\n\n".join(paras)
        for title, paras in sections.items() if paras
    }
    return chunks, page_texts


def process_docx(path, filename):
    document = Document(path)
    core = document.core_properties

    title = (core.title or "").strip()
    if not title:
        for para in document.paragraphs:
            if para.text.strip():
                title = para.text.strip()
                break
    if not title:
        title = filename

    author = (core.author or "").strip() or None

    chunks, page_texts = read_docx(path, filename)
    file_meta = {
        "filename": filename,
        "file_type": "docx",
        "guessed_title": title,
        "author": author,
        "extent_label": f"{len(page_texts)} section(s)"
    }
    return chunks, page_texts, file_meta


# ------------------------------------------------------------------
# 2c. Unified processing across all supported file types
# ------------------------------------------------------------------
def process_all_documents(folder):
    all_chunks = []
    all_page_texts = {}
    metadata = {"files": []}

    for f in list_documents(folder):
        path = os.path.join(folder, f)
        ext = os.path.splitext(f)[1].lower()

        if ext == ".pdf":
            chunks, page_texts, file_meta = process_pdf(path, f)
        elif ext == ".docx":
            chunks, page_texts, file_meta = process_docx(path, f)
        else:
            continue  # unsupported, shouldn't happen given list_documents filter

        metadata["files"].append(file_meta)
        all_chunks.extend(chunks)
        all_page_texts[f] = page_texts

    return all_chunks, metadata, all_page_texts


# ------------------------------------------------------------------
# 3. Cache load/save
# ------------------------------------------------------------------
def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as fh:
            return json.load(fh)
    return None


def save_cache(fingerprint, chunks, metadata, page_texts):
    with open(CACHE_FILE, "w", encoding="utf-8") as fh:
        json.dump({
            "fingerprint": fingerprint,
            "version": CACHE_VERSION,
            "chunks": chunks,
            "metadata": metadata,
            "page_texts": page_texts
        }, fh)


def get_chunks_and_metadata():
    """Returns (chunks, metadata, page_texts), using the on-disk cache
    whenever the documents/ folder hasn't changed since last run."""
    fingerprint = folder_fingerprint(DOCS_DIR)
    cached = load_cache()
    if cached and cached.get("fingerprint") == fingerprint and cached.get("version") == CACHE_VERSION:
        return cached["chunks"], cached["metadata"], cached.get("page_texts", {})

    chunks, metadata, page_texts = process_all_documents(DOCS_DIR)
    if chunks:
        save_cache(fingerprint, chunks, metadata, page_texts)
    return chunks, metadata, page_texts


# ------------------------------------------------------------------
# 4. Retrieval: chunk-level, then sentence-level for concise snippets
# ------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def build_vectorizer(chunks_tuple):
    """chunks_tuple is a tuple of texts (hashable) so st.cache_resource
    can key on it correctly."""
    texts = list(chunks_tuple)
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    vectors = vectorizer.fit_transform(texts)
    return vectorizer, vectors


def display_name(filename, metadata):
    """Prefer the guessed title (e.g. 'Originals') over the raw filename
    for citations."""
    for f in metadata["files"]:
        if f["filename"] == filename:
            title = f.get("guessed_title")
            return title if title else filename
    return filename


def get_file_type(filename, metadata):
    for f in metadata["files"]:
        if f["filename"] == filename:
            return f.get("file_type", "pdf")
    return "pdf"


def relevance_badge(score):
    """Raw TF-IDF cosine scores (often 0.1-0.3) are meaningless to most
    users, so translate them into a plain-language signal instead."""
    if score >= 0.30:
        return "🟢 Strong match"
    elif score >= 0.15:
        return "🟡 Moderate match"
    return "🟠 Loose match"


def split_sentences(text):
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if len(s.strip()) > 15]


DEFINITION_TERM_PATTERN = re.compile(
    r"^(?:what\s+(?:is|are)|define|what's|meaning\s+of)\s+(.+)", re.IGNORECASE
)


def extract_definition_term(question):
    match = DEFINITION_TERM_PATTERN.match(question.strip().rstrip("?. "))
    if match:
        return match.group(1).strip()
    return None


def concise_snippet(chunk_text, question, term):
    """Pick the 1-2 best sentences within a chunk for a short answer,
    boosting sentences that look like a definition when relevant."""
    sentences = split_sentences(chunk_text)
    if not sentences:
        return chunk_text[:400]

    sent_vectorizer = TfidfVectorizer(stop_words="english")
    try:
        sent_vectors = sent_vectorizer.fit_transform(sentences)
        q_vec = sent_vectorizer.transform([question])
        sent_scores = cosine_similarity(q_vec, sent_vectors)[0]

        if term:
            def_pattern = re.compile(
                rf"\b{re.escape(term.lower())}\b\s+(is|are|refers to|means)\b",
                re.IGNORECASE
            )
            for idx, sent in enumerate(sentences):
                if def_pattern.search(sent):
                    sent_scores[idx] += 1.0

        top_idx = sorted(sent_scores.argsort()[::-1][:MAX_ANSWER_SENTENCES])
        return " ".join(sentences[i] for i in top_idx)
    except ValueError:
        return " ".join(sentences[:MAX_ANSWER_SENTENCES])


def get_top_matches(chunks, question, top_n=TOP_MATCHES):
    """Returns up to top_n distinct matches, each with a short snippet plus
    the full chunk text (for the 'read more' expander)."""
    texts = tuple(c["text"] for c in chunks)
    vectorizer, vectors = build_vectorizer(texts)
    question_vector = vectorizer.transform([question])
    scores = cosine_similarity(question_vector, vectors)[0]
    ranked = scores.argsort()[::-1]

    term = extract_definition_term(question)
    matches = []
    seen_locations = set()

    for i in ranked:
        if scores[i] <= 0 or len(matches) >= top_n:
            break
        chunk = chunks[i]
        key = (chunk["source"], chunk["location"])
        if key in seen_locations:
            continue  # skip duplicate paragraphs from the same page/section
        seen_locations.add(key)

        matches.append({
            "source": chunk["source"],
            "location": chunk["location"],
            "snippet": concise_snippet(chunk["text"], question, term),
            "full_text": chunk["text"],
            "score": round(float(scores[i]), 2)
        })

    return matches


# ------------------------------------------------------------------
# 4b. Conversational summary — still 100% extractive (no LLM), just
# stitched into flowing prose instead of a bare list of quotes.
# ------------------------------------------------------------------
FOLLOWUP_CONNECTORS = [
    "I also found this in",
    "Additionally,",
    "It's also worth noting from",
]


def build_conversational_summary(matches, metadata_):
    if not matches:
        return None

    lead_title = display_name(matches[0]["source"], metadata_)
    lines = [f"Here's what I found — **{lead_title}** ({matches[0]['location']}) says:\n\"{matches[0]['snippet']}\""]

    for i, m in enumerate(matches[1:], start=1):
        title = display_name(m["source"], metadata_)
        connector = FOLLOWUP_CONNECTORS[(i - 1) % len(FOLLOWUP_CONNECTORS)]
        lines.append(f"{connector} **{title}** ({m['location']}):\n\"{m['snippet']}\"")

    return "\n\n".join(lines)


# ------------------------------------------------------------------
# 4c. Q&A logging — every question and answer gets appended to a CSV
# ------------------------------------------------------------------
def log_interaction(question, answer_text, matches, metadata_, username):
    timestamp = datetime.now().isoformat(timespec="seconds")
    if matches:
        sources = "; ".join(
            f"{display_name(m['source'], metadata_)} ({m['location']}, score {m['score']})"
            for m in matches
        )
    else:
        sources = ""

    file_exists = os.path.exists(QA_LOG_FILE)
    try:
        with open(QA_LOG_FILE, "a", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            if not file_exists:
                writer.writerow(["timestamp", "username", "question", "answer", "sources"])
            writer.writerow([timestamp, username, question, answer_text, sources])
    except OSError:
        pass  # don't let logging failures break the chat experience


# ------------------------------------------------------------------
# 4d. Authentication — file-based, interim until SSO is wired up
# ------------------------------------------------------------------
def load_users():
    if not os.path.exists(USERS_FILE):
        return {}
    with open(USERS_FILE, "r", encoding="utf-8") as fh:
        return json.load(fh)


def save_users(users):
    with open(USERS_FILE, "w", encoding="utf-8") as fh:
        json.dump(users, fh, indent=2)


def hash_password(password, salt_hex=None):
    """PBKDF2-HMAC-SHA256 with a per-user random salt — no plaintext
    passwords ever touch disk, even in this interim setup."""
    salt = bytes.fromhex(salt_hex) if salt_hex else os.urandom(16)
    hashed = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 100_000)
    return salt.hex(), hashed.hex()


def verify_password(password, salt_hex, expected_hash_hex):
    _, computed_hash_hex = hash_password(password, salt_hex)
    return hmac.compare_digest(computed_hash_hex, expected_hash_hex)


def check_login(username, password):
    users = load_users()
    record = users.get(username)
    if not record:
        return False
    return verify_password(password, record["salt"], record["hash"])


def register_user(username, password):
    """Self-service account creation from within the app. Returns
    (success, message). Same hashing scheme as manage_users.py, so
    accounts created either way are fully interchangeable."""
    username = username.strip()
    if not username:
        return False, "Username cannot be empty."
    if len(password) < 6:
        return False, "Password must be at least 6 characters."

    users = load_users()
    if username in users:
        return False, f"Username '{username}' is already taken."

    salt_hex, hash_hex = hash_password(password)
    users[username] = {"salt": salt_hex, "hash": hash_hex}
    save_users(users)
    return True, f"Account '{username}' created. Switch to the Sign in tab to log in."


def render_login():
    st.title(f"🔐 {BOT_NAME} — Sign in")

    tab_signin, tab_signup = st.tabs(["Sign in", "Create account"])

    with tab_signin:
        if not load_users():
            st.info("No accounts exist yet — use the **Create account** tab to make one.")

        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Sign in")

        if submitted:
            if check_login(username, password):
                st.session_state.authenticated = True
                st.session_state.username = username
                st.rerun()
            else:
                st.error("Invalid username or password.")

    with tab_signup:
        st.caption("Create your own login — no admin needed for this interim setup.")
        with st.form("signup_form"):
            new_username = st.text_input("Choose a username", key="signup_username")
            new_password = st.text_input("Choose a password", type="password", key="signup_password")
            confirm_password = st.text_input("Confirm password", type="password", key="signup_confirm")
            signup_submitted = st.form_submit_button("Create account")

        if signup_submitted:
            if new_password != confirm_password:
                st.error("Passwords don't match.")
            else:
                success, message = register_user(new_username, new_password)
                if success:
                    st.success(message)
                else:
                    st.error(message)


# ------------------------------------------------------------------
# 4e. Session logging — one row per login, now tied to a real username
# since we have actual accounts instead of anonymous visits.
# ------------------------------------------------------------------
def log_session_once():
    if st.session_state.get("session_logged"):
        return
    st.session_state.session_logged = True
    username = st.session_state.get("username", "unknown")
    timestamp = datetime.now().isoformat(timespec="seconds")
    file_exists = os.path.exists(SESSION_LOG_FILE)
    try:
        with open(SESSION_LOG_FILE, "a", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            if not file_exists:
                writer.writerow(["timestamp", "username"])
            writer.writerow([timestamp, username])
    except OSError:
        pass


# ------------------------------------------------------------------
# 4f. Statistics page — FAQs, user/session counts, simple charts
# ------------------------------------------------------------------
def normalize_question(q):
    q = re.sub(r"\s+", " ", q.strip().lower())
    return q.rstrip("!?. ")


def load_qa_log_df():
    if not os.path.exists(QA_LOG_FILE):
        return pd.DataFrame(columns=["timestamp", "username", "question", "answer", "sources"])
    return pd.read_csv(QA_LOG_FILE)


def load_session_log_df():
    if not os.path.exists(SESSION_LOG_FILE):
        return pd.DataFrame(columns=["timestamp", "username"])
    return pd.read_csv(SESSION_LOG_FILE)


def render_stats_page():
    st.header("📊 Statistics")
    qa_df = load_qa_log_df()
    session_df = load_session_log_df()
    distinct_users = session_df["username"].nunique() if not session_df.empty else 0

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total questions asked", len(qa_df))
    col2.metric("Total logins", len(session_df))
    col3.metric("Distinct users", distinct_users)
    doc_answered = int((qa_df["sources"].fillna("").astype(str).str.strip() != "").sum()) if not qa_df.empty else 0
    col4.metric("Answered from documents", doc_answered)

    if not session_df.empty:
        st.subheader("👥 Logins by User")
        logins_by_user = session_df["username"].value_counts()
        st.bar_chart(logins_by_user)

    if qa_df.empty:
        st.info("No questions logged yet — ask something in the 💬 Chat tab first!")
        return

    st.subheader("🔥 Frequently Asked Questions")
    qa_df["normalized"] = qa_df["question"].astype(str).apply(normalize_question)
    top_questions = qa_df["normalized"].value_counts().head(10)
    top_df = top_questions.rename_axis("question").reset_index(name="count")
    st.bar_chart(top_df.set_index("question")["count"])
    st.dataframe(top_df, use_container_width=True, hide_index=True)

    st.subheader("📅 Questions Over Time")
    qa_df["date"] = pd.to_datetime(qa_df["timestamp"], errors="coerce").dt.date
    per_day = qa_df.dropna(subset=["date"]).groupby("date").size()
    if len(per_day) > 1:
        st.line_chart(per_day)
    else:
        st.caption("Not enough days of data yet for a trend line.")

    st.subheader("🧭 Answer Source Breakdown")
    breakdown = pd.Series({
        "Answered from documents": doc_answered,
        "Generic / small talk": len(qa_df) - doc_answered
    })
    st.bar_chart(breakdown)


# ------------------------------------------------------------------
# 5. Generic / small-talk handling (no retrieval needed)
# ------------------------------------------------------------------
GREETINGS = {"hi", "hello", "hey", "hii", "hiya", "good morning", "good afternoon", "good evening"}
THANKS = {"thanks", "thank you", "thx", "ty"}
BYE = {"bye", "goodbye", "see you", "exit", "quit"}


def handle_generic(question, metadata):
    q = question.strip().lower().rstrip("!?. ")

    if re.search(r"\byour name\b", q) or q in {"who are you", "what are you"}:
        return f"I'm **{BOT_NAME}** — your local document Q&A assistant. Ask me anything about the files loaded from the documents/ folder!"

    if q in GREETINGS or any(q.startswith(g) for g in GREETINGS):
        return f"Hi! I'm {BOT_NAME}. Ask me anything about the document(s) loaded from the documents/ folder."

    if any(t in q for t in THANKS):
        return "You're welcome! Anything else you'd like to ask?"

    if q in BYE:
        return "Goodbye! Come back anytime you have more questions."

    if "how are you" in q:
        return f"Doing well, thanks for asking! I'm {BOT_NAME} — what can I help you find in the document?"

    if re.search(r"\bauthor\b", q) or re.search(r"who\s+(wrote|is the author)", q):
        if not metadata["files"]:
            return "I don't have any documents loaded yet — add one to the documents/ folder."
        lines = []
        for f in metadata["files"]:
            author = f.get("author") or "not stated in the document"
            lines.append(f"- **{f['guessed_title']}** — author: {author}")
        return "\n".join(lines)

    if re.search(r"\b(title|name of the (book|document|pdf|file))\b", q):
        if metadata["files"]:
            lines = [f"- **{f['guessed_title']}** ({f['filename']})" for f in metadata["files"]]
            return "Here's what I have loaded:\n" + "\n".join(lines)
        return "I don't have any documents loaded yet — add one to the documents/ folder."

    if re.search(r"how many (pages|documents|files|pdfs|sections)", q):
        lines = [f"- **{f['guessed_title']}**: {f['extent_label']}" for f in metadata["files"]]
        return f"I have {len(metadata['files'])} document(s) loaded:\n" + "\n".join(lines)

    # Broad, typo-tolerant match for "what documents/files do you have" style
    # questions — keyword-based rather than a strict phrase match, since real
    # users phrase this many different ways ("what are you documents you have").
    if re.search(r"\b(document|doc|pdf|file)s?\b", q) and re.search(r"\b(have|got|loaded|available)\b", q):
        if metadata["files"]:
            names = ", ".join(f["filename"] for f in metadata["files"])
            return f"Loaded documents: {names}"
        return "No documents are loaded yet."

    return None  # not a generic question -> fall through to retrieval


# ------------------------------------------------------------------
# 6. Streamlit chat UI
# ------------------------------------------------------------------
st.set_page_config(page_title=BOT_NAME, page_icon="🎓")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    render_login()
    st.stop()

log_session_once()

os.makedirs(DOCS_DIR, exist_ok=True)
chunks, metadata, page_texts = get_chunks_and_metadata()

if not chunks:
    st.title(f"🎓 {BOT_NAME}")
    st.warning(f"No PDF or Word files found in `{DOCS_DIR}`. Add files there and refresh the page.")
    st.stop()

with st.sidebar:
    st.caption(f"Signed in as **{st.session_state.username}**")
    if st.button("🚪 Log out"):
        for key in ("authenticated", "username", "session_logged", "messages"):
            st.session_state.pop(key, None)
        st.rerun()
    st.divider()

    page = st.radio("Navigate", ["💬 Chat", "📊 Statistics"], label_visibility="collapsed")
    st.divider()

    st.subheader("Loaded documents")
    for f in metadata["files"]:
        st.write(f"**{f['guessed_title']}** ({f['file_type'].upper()}, {f['extent_label']})")
    if st.button("🔄 Re-scan documents/ folder"):
        st.cache_resource.clear()
        if os.path.exists(CACHE_FILE):
            os.remove(CACHE_FILE)
        st.rerun()

    st.divider()
    st.subheader("Q&A log")
    if os.path.exists(QA_LOG_FILE):
        with open(QA_LOG_FILE, "rb") as fh:
            st.download_button("⬇️ Download qa_log.csv", fh, file_name="qa_log.csv", mime="text/csv")
    else:
        st.caption("No questions logged yet.")

if page == "📊 Statistics":
    render_stats_page()
    st.stop()

st.title(f"🎓 {BOT_NAME}")

if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "type": "text",
         "content": f"Hi! I'm {BOT_NAME}. Ask me anything about the document(s) loaded from the documents/ folder."}
    ]


def render_answer(summary, matches, metadata_, page_texts_, msg_idx):
    """Shows the conversational summary up top (the natural-reading part),
    then the individual match cards underneath as verifiable evidence —
    so the flowing answer and the raw sources stay visually distinct."""
    st.markdown(summary)
    st.divider()
    st.caption(f"📚 Sources — {len(matches)} passage(s) behind this answer:")

    for rank, m in enumerate(matches, start=1):
        title = display_name(m["source"], metadata_)
        icon = "📄" if get_file_type(m["source"], metadata_) == "pdf" else "📝"
        badge = relevance_badge(m["score"])
        heading = "✅ Best match" if rank == 1 else f"Match {rank}"

        with st.container(border=True):
            header_col, badge_col = st.columns([2, 1])
            with header_col:
                st.markdown(f"**{heading}**")
            with badge_col:
                st.markdown(badge)

            # Blockquote styling signals "this is quoted from the source",
            # not GurucoolBOT's own generated words.
            st.markdown(f"> {m['snippet']}")

            st.markdown(f"{icon} **{title}** · {m['location']}")

            full_text = page_texts_.get(m["source"], {}).get(m["location"], m["full_text"])
            loc_label = "page" if m["location"].startswith("Page") else "section"
            with st.expander(f"🔍 View full {loc_label} (ref #{msg_idx}.{rank})"):
                st.caption(f"Raw relevance score: {m['score']}")
                st.write(full_text)


# Replay chat history
for idx, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        if msg["type"] == "text":
            st.markdown(msg["content"])
        elif msg["type"] == "matches":
            render_answer(msg["summary"], msg["matches"], metadata, page_texts, idx)

# New input
question = st.chat_input("Ask a question...")
if question:
    st.session_state.messages.append({"role": "user", "type": "text", "content": question})
    with st.chat_message("user"):
        st.markdown(question)

    generic_reply = handle_generic(question, metadata)

    with st.chat_message("assistant"):
        if generic_reply is not None:
            st.markdown(generic_reply)
            st.session_state.messages.append({"role": "assistant", "type": "text", "content": generic_reply})
            log_interaction(question, generic_reply, [], metadata, st.session_state.username)
        else:
            matches = get_top_matches(chunks, question)
            if not matches:
                reply = "Sorry, I couldn't find anything relevant to that in the document(s)."
                st.markdown(reply)
                st.session_state.messages.append({"role": "assistant", "type": "text", "content": reply})
                log_interaction(question, reply, [], metadata, st.session_state.username)
            else:
                summary = build_conversational_summary(matches, metadata)
                new_idx = len(st.session_state.messages)
                render_answer(summary, matches, metadata, page_texts, new_idx)
                st.session_state.messages.append(
                    {"role": "assistant", "type": "matches", "summary": summary, "matches": matches}
                )
                log_interaction(question, summary, matches, metadata, st.session_state.username)
